# docx_generator.py
import streamlit as st
import requests
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def _add_hyperlink(paragraph, text, url):
    """Adiciona um hyperlink a um parágrafo no documento."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
    rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def _draw_report_content_docx(document, info: dict, product_url: str):
    """Desenha o conteúdo do relatório de UM produto no documento Word."""
    # Bloco de Informações do Produto
    p_title = document.add_heading(info.get('product_title', 'Título não encontrado'), level=1)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_asin = document.add_paragraph()
    p_asin.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_asin.add_run('ASIN: ').bold = True
    p_asin.add_run(info.get('asin', 'N/A'))
    
    # Adiciona o link do produto de forma clicável
    p_link = document.add_paragraph()
    p_link.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_link.add_run('Link do Produto: ').bold = True
    _add_hyperlink(p_link, product_url, product_url)

    # Bloco de Análise de Inconsistências
    p_header_report = document.add_heading("Relatório de Inconsistências e Melhorias", level=2)
    p_header_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    report_text = info.get('report', 'Nenhum relatório disponível.')
    p_report = document.add_paragraph()
    # Processa o negrito do Markdown
    parts = re.split(r'(\*\*.*?\*\*)', report_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p_report.add_run(part[2:-2]).bold = True
        else:
            p_report.add_run(part)

    # Bloco de Imagens do Produto
    document.add_heading("Imagens do Produto", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_urls = info.get('product_photos', [])
    if not image_urls:
        document.add_paragraph("Nenhuma imagem adicional encontrada.")
    
    for i, url in enumerate(image_urls):
        try:
            response = requests.get(url, timeout=20)
            response.raise_for_status()
            image_stream = io.BytesIO(response.content)
            # Adiciona a imagem centralizada e com tamanho reduzido
            p_img = document.add_picture(image_stream, width=Inches(5.0))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Adiciona a legenda
            p_caption = document.add_paragraph(f"Imagem {i+1}")
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            document.add_paragraph(f"Erro ao carregar imagem {i+1}.").alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"Erro ao baixar imagem para DOCX: {e}")

def create_single_docx_report(info: dict, product_url: str):
    """Cria e retorna um DOCX para um único relatório."""
    document = Document()

    # Adiciona o branding
    logo_path = 'globald_logo_512x512_original.jpg'
    if os.path.exists(logo_path):
        p_logo = document.add_picture(logo_path, width=Inches(1.5))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        st.warning(f"Arquivo do logo '{logo_path}' não encontrado.")
    
    p_tagline = document.add_paragraph('AI Compliance Relatório by www.GlobalD.ai')
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    _draw_report_content_docx(document, info, product_url)
    
    # Salva o documento em um buffer de memória
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_batch_docx_report(batch_results: list, urls: list):
    """Cria um DOCX consolidado a partir de uma lista de resultados."""
    document = Document()

    # Adiciona o branding apenas uma vez, na primeira página
    logo_path = 'globald_logo_512x512_original.jpg'
    if os.path.exists(logo_path):
        p_logo = document.add_picture(logo_path, width=Inches(1.5))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_tagline = document.add_paragraph('AI Compliance Relatório by www.GlobalD.ai')
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Espaçamento

    for i, result_info in enumerate(batch_results):
        product_url = urls[i] if i < len(urls) else "URL não encontrada"
        
        # Adiciona uma quebra de página antes de cada novo produto (exceto o primeiro)
        if i > 0:
            document.add_page_break()
            
        _draw_report_content_docx(document, result_info, product_url)

    # Salva o documento em um buffer de memória
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io